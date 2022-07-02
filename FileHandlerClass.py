# -*- coding: utf-8 -*-
"""
Created on Fri Aug  6 21:21:46 2021

@author: Lina, jhowl
"""
import pandas as pd                               
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import os


class FileHandler:
"""
Обработка результатов теста и предоставление результатов в удобном для чтения виде 
"""
  
    def __init__(self, filename):
        self.df = pd.read_csv(filename)
        
    def ReadFile(self):
    """
    Изменение формы таблицы
    """
        self.df['Балл ситуативной тревожности'] = ''
        self.df['Ситуативная тревожность'] = ''
        self.df['Балл личностной тревожности'] = ''
        self.df['Личностная тревожность'] = ''
        self.df_width = self.df.shape[1]
        self.df_height = self.df.shape[0]
    
    def Handler(self):
    """
    Обработка результатов теста
    """
        for j in range(3,self.df_width):
            for i in range(self.df_height):
                if (self.df.iloc[i,j]) == 'Нет, это совсем не так.':
                    self.df.iloc[i,j] = '1'
                if (self.df.iloc[i,j]) == 'Пожалуй, так.':
                    self.df.iloc[i,j] = '2'
                if (self.df.iloc[i,j]) == 'Верно.':
                    self.df.iloc[i,j] = '3'
                if (self.df.iloc[i,j]) == 'Совершенно верно.':
                    self.df.iloc[i,j] = '4'

        for j in range(3,self.df_width):
            for i in range(self.df_height):
                if (self.df.iloc[i,j]) == 'Почти никогда.':
                    self.df.iloc[i,j] = '1'
                if (self.df.iloc[i,j]) == 'Иногда.':
                    self.df.iloc[i,j] = '2'
                if (self.df.iloc[i,j]) == 'Часто.':
                    self.df.iloc[i,j] = '3'
                if (self.df.iloc[i,j]) == 'Почти всегда.':
                    self.df.iloc[i,j] = '4' 

        sit_anx1 = [5, 6, 8, 9, 11, 14, 15, 16, 19, 20]
        sit_anx2 = [3, 4, 7, 10, 12, 13, 17, 18, 21, 22]
        
        pers_anx1 = [24,25, 26, 27, 30, 31, 33, 34, 36, 37, 39, 40, 42]
        pers_anx2 = [23, 28, 29, 32, 35, 38, 41]
        
        var_sit_anx1 = 0
        sum_sit_anx1 = [0]* self.df_height
        var_sit_anx2 = 0
        sum_sit_anx2 = [0]* self.df_height
        var_pers_anx1 = 0
        sum_pers_anx1 = [0]* self.df_height
        var_pers_anx2 = 0
        sum_pers_anx2 = [0]* self.df_height
        
        for i in range(self.df_height):
            for j in sit_anx1:
                var_sit_anx1 = var_sit_anx1 + int(self.df.iloc[i,j])
            sum_sit_anx1[i] = var_sit_anx1
            var_sit_anx1 = 0
        sum_sit_anx1
        
        for i in range(self.df_height):
            for j in sit_anx2:
                var_sit_anx2 = var_sit_anx2 + int(self.df.iloc[i,j])
            sum_sit_anx2[i] = var_sit_anx2
            var_sit_anx2 = 0
        sum_sit_anx2
        
        for i in range(self.df_height):
            for j in pers_anx1:
                var_pers_anx1 = var_pers_anx1 + int(self.df.iloc[i,j])
            sum_pers_anx1[i] = var_pers_anx1
            var_pers_anx1 = 0
        sum_pers_anx1
        
        for i in range(self.df_height):
            for j in pers_anx2:
                var_pers_anx2 = var_pers_anx2 + int(self.df.iloc[i,j])
            sum_pers_anx2[i] = var_pers_anx2
            var_pers_anx2 = 0
        sum_pers_anx2
        
        all_sit_anx = [sum_sit_anx1[i] - sum_sit_anx2[i] + 35 for i in range(self.df_height)]
        all_pers_anx = [sum_pers_anx1[i] - sum_pers_anx2[i] + 35 for i in range(self.df_height)]
        
        self.df['Балл ситуативной тревожности'] = all_sit_anx
        self.df['Балл личностной тревожности'] = all_pers_anx
        
        for i in range(self.df_height): 
            if 20 <= (self.df.loc[i,'Балл ситуативной тревожности']) <= 30:
                self.df.loc[i,'Ситуативная тревожность'] = 'низкий уровень тревожности'
            if 31 <= (self.df.loc[i,'Балл ситуативной тревожности']) <= 45:
                self.df.loc[i,'Ситуативная тревожность'] = 'средний уровень тревожности'
            if 46 < (self.df.loc[i,'Балл ситуативной тревожности']):
                self.df.loc[i,'Ситуативная тревожность'] = 'высокий уровень тревожности'
                
        for i in range(self.df_height): 
            if 20 <= (self.df.loc[i,'Балл личностной тревожности']) <= 30:
                self.df.loc[i,'Личностная тревожность'] = 'низкий уровень тревожности'
            if 31 <= (self.df.loc[i,'Балл личностной тревожности']) <= 45:
                self.df.loc[i,'Личностная тревожность'] = 'средний уровень тревожности'
            if 46 < (self.df.loc[i,'Балл личностной тревожности']):
                self.df.loc[i,'Личностная тревожность'] = 'высокий уровень тревожности'
     
        self.df_f = self.df[['ФИО', 'Балл ситуативной тревожности', 'Ситуативная тревожность', 'Балл личностной тревожности', 'Личностная тревожность']]
     
        low_pers_anx = 0
        med_pers_anx = 0
        high_pers_anx = 0
        self.marks_list_1 = ['Низкий уровень личностной тревожности имеют', 'Средний уровень личностной тревожности имеют', 'Высокий уровень личностной тревожности имеют']
        self.marks_list_2 = ['Низкий уровень ситуативной тревожности имеют', 'Средний уровень ситуативной тревожности имеют', 'Высокий уровень ситуативной тревожности имеют']
        self.sit_percent = [0]*3
        self.pers_percent = [0]*3
        
        
        for i in range(self.df_height):
            if self.df.loc[i,'Личностная тревожность'] == 'низкий уровень тревожности':
                low_pers_anx = low_pers_anx+1
            if self.df.loc[i,'Личностная тревожность'] == 'средний уровень тревожности':
                med_pers_anx = med_pers_anx+1
            if self.df.loc[i,'Личностная тревожность'] == 'высокий уровень тревожности':
                high_pers_anx = high_pers_anx+1
        
        pers_anx_list = [low_pers_anx,med_pers_anx,high_pers_anx]
        
        j = 0
        for i in pers_anx_list:
            self.pers_percent[j] = (i*100)/(low_pers_anx+med_pers_anx+high_pers_anx)
            print(self.marks_list_1[j], self.pers_percent[j], '% тестируемых')
            j += 1
        
        low_sit_anx = 0
        med_sit_anx = 0
        high_sit_anx = 0
        

        
        for i in range(self.df_height):
            if self.df.loc[i,'Ситуативная тревожность'] == 'низкий уровень тревожности':
                low_sit_anx = low_sit_anx+1
            if self.df.loc[i,'Ситуативная тревожность'] == 'средний уровень тревожности':
                med_sit_anx = med_sit_anx+1
            if self.df.loc[i,'Ситуативная тревожность'] == 'высокий уровень тревожности':
                high_sit_anx = high_sit_anx+1
        
        sit_anx_list = [low_sit_anx,med_sit_anx,high_sit_anx]
        
        j = 0
        for i in sit_anx_list:
            self.sit_percent[j] = (i*100)/(low_sit_anx+med_sit_anx+high_sit_anx)
            print(self.marks_list_2[j], self.sit_percent[j], '% тестируемых')
            j += 1
          
        self.df_f2 = self.df_f.copy()
        self.df_f2['Интерпретация результатов сит. трев'] = ''
        self.df_f2['Интерпретация результатов личн. трев'] = ''  
            
        low_anx_text = "Низкие значения уровня тревожности свидетельствуют о сниженном чувстве ответственности\
        и необходимости обратить внимание на мотивы деятельности, выполняемой человеком. В некоторых\
        случаях низкая тревожность в показателях теста является результатом активного вытеснения\
        личностью высокой тревоги с целью показать себя«социально желательным."  
        
        high_anx_text = "Высокие значения уровня тревожности предполагают склонность к появлению"\
        +"состояния тревоги у человека в ситуациях оценки его компетентности и свидетельствуют о"\
        +"необходимости снизить субъективную значимость ситуации, перенести акцент на"\
        +"осмысление деятельности. Высокая личностная тревожности характеризуются также"\
        +"устойчивой склонностью воспринимать большой круг ситуаций как угрожающие и"\
        +"коррелирует с эмоциональными и невротическими срывами."\
        +"Высокое ситуативное состояние тревоги характеризуется напряжением,"\
        +"беспокойством, нервозностью. Это вызывает нарушение внимания, нарушение тонкой"\
        +"координации."
        
        self.df_f3 = self.df_f2.copy()
        self.df_f3['Интерпретация результатов сит. трев'] = self.df_f3['Ситуативная тревожность'].apply(lambda r: low_anx_text if r == 'низкий уровень тревожности' else ('norm' if r == 'средний уровень тревожности' else high_anx_text))
        self.df_f3['Интерпретация результатов личн. трев'] = self.df_f3['Личностная тревожность'].apply(lambda r: low_anx_text if r == 'низкий уровень тревожности' else ('norm' if r == 'средний уровень тревожности' else high_anx_text))
        
    def WriteFile(self, path):
    """
    Запись результатов обработки в новый файл
    """
        print(path)
        doc = docx.Document() 
        styles = doc.styles
        new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 1']
        font = new_heading_style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor.from_string('000000')
        par_style = styles.add_style('par', WD_STYLE_TYPE.PARAGRAPH)
        par_style.base_style = styles['Body Text']
        par_font = par_style.font
        par_font.name = 'Times New Roman'
        par_font.size = Pt(14)
        table = doc.add_table(self.df_f.shape[0]+1, self.df_f.shape[1])
        
        table.style = 'Table Grid'
        # add the header rows.
        for j in range(self.df_f.shape[-1]):
            table.cell(0,j).text = self.df_f.columns[j]
        
        # add the rest of the data frame
        for i in range(self.df_f.shape[0]):
            for j in range(self.df_f.shape[-1]):
                table.cell(i+1,j).text = str(self.df_f.values[i,j])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size= Pt(14)
                        paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for cell in table.rows[0].cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True
        doc.add_paragraph('Результат тестирования по группе ' + str(self.df.loc[1,'Группа']), style='New Heading')
        for i in range(0,3):
            doc.add_paragraph(str(self.marks_list_1[i]) + ' ' + str(self.pers_percent[i]) + ' ' + '% тестируемых', style='par')
        for i in range(0,3):
            doc.add_paragraph(str(self.marks_list_2[i]) + ' ' + str(self.sit_percent[i]) + ' ' + '% тестируемых', style='par')
        if not os.path.exists(path+'/Результат'):
            os.mkdir(path+'/Результат')
        doc.save(path+"/Результат/Сводная таблица группы "+str(self.df.loc[1,'Группа'])+".docx")
        
        pers_doc = docx.Document()
        pers_styles = pers_doc.styles
        new_heading_style = pers_styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = pers_styles['Heading 1']
        font = new_heading_style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor.from_string('000000')
        para_style = pers_styles.add_style('para', WD_STYLE_TYPE.PARAGRAPH)
        para_style.base_style = pers_styles['Body Text']
        para_font = para_style.font
        para_font.name = 'Times New Roman'
        para_font.size = Pt(14)
        pers_doc.add_paragraph('Индивидуальные результаты тестирования', style='New Heading')
        for i in range(self.df_f3.shape[0]):
            pers_doc.add_paragraph(str(self.df_f3.loc[i, 'ФИО']), style='New Heading')
            pers_doc.add_paragraph('Ситуативная тревожность:', style='para')
            pers_doc.add_paragraph(str(self.df_f3.loc[i, 'Интерпретация результатов сит. трев']), style='para')
            pers_doc.add_paragraph('Личностная тревожность:', style='para')
            pers_doc.add_paragraph(str(self.df_f3.loc[i, 'Интерпретация результатов личн. трев']), style='para')
            pers_doc.add_page_break()
        if not os.path.exists(path+'/Результат'):
            os.mkdir(path+'/Результат')
        pers_doc.save(path+"/Результат/Индивидуальные результаты.docx")
